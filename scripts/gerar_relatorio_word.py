"""
Gera o Relatório Final do Projeto OnboardDev em formato Word (.docx)
Disciplina: Ferramentas para Geração de Ideias — UFRJ Politécnica
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================
# Configuração
# ============================================================
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Relatorio_Final_OnboardDev.docx")

TEAM = [
    "Bruno Levorato",
    "Rafael Ribeiro",
    "Renato Gondin",
    "Vitor Mello",
]
PROFESSOR = "Juliana França"
UNIVERSITY = "Universidade Federal do Rio de Janeiro — UFRJ"
SCHOOL = "Escola Politécnica"
COURSE = "Ferramentas para Geração de Ideias"
DATE = "Rio de Janeiro, Junho de 2026"

# Cores
PRIMARY_COLOR = RGBColor(0x63, 0x66, 0xF1)  # Indigo
DARK_COLOR = RGBColor(0x1E, 0x1E, 0x2E)
GRAY_COLOR = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
TABLE_HEADER_BG = "4338CA"
TABLE_ALT_BG = "EEF2FF"


def set_cell_shading(cell, color_hex):
    """Define a cor de fundo de uma célula de tabela."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """Adiciona bordas sutis a uma tabela."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "D1D5DB")
        borders.append(border)
    tblPr.append(borders)


def create_styled_table(doc, headers, rows, col_widths=None):
    """Cria uma tabela estilizada com headers coloridos."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_heading_styled(doc, text, level=1):
    """Adiciona heading com estilo customizado."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK_COLOR
    return heading


def add_paragraph_styled(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    """Adiciona parágrafo com estilo."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Adiciona item de lista."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


# ============================================================
# Geração do Documento
# ============================================================
def generate_report():
    doc = Document()

    # ---- Configurar estilos padrão ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BLACK

    # Configurar margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ================================================================
    # CAPA
    # ================================================================
    for _ in range(4):
        doc.add_paragraph("")

    add_paragraph_styled(doc, UNIVERSITY, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_styled(doc, SCHOOL, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_styled(doc, f"Disciplina: {COURSE}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY_COLOR)
    add_paragraph_styled(doc, f"Professora: {PROFESSOR}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY_COLOR)

    doc.add_paragraph("")
    doc.add_paragraph("")

    add_paragraph_styled(doc, "OnboardDev", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY_COLOR)
    add_paragraph_styled(
        doc,
        "Plataforma Inteligente de Onboarding para Desenvolvedores",
        bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER
    )

    doc.add_paragraph("")
    add_paragraph_styled(
        doc,
        "Ideação e Prototipagem de Soluções Inovadoras para a Indústria de Software",
        italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY_COLOR
    )

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    add_paragraph_styled(doc, "Integrantes:", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name in TEAM:
        add_paragraph_styled(doc, name, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("")
    doc.add_paragraph("")
    add_paragraph_styled(doc, DATE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY_COLOR)

    doc.add_page_break()

    # ================================================================
    # SUMÁRIO (placeholder manual)
    # ================================================================
    add_heading_styled(doc, "Sumário", level=1)
    toc_items = [
        ("1.", "Introdução", "3"),
        ("2.", "Identificação do Problema", "4"),
        ("3.", "Pesquisa e Empatia", "5"),
        ("4.", "Sessão de Ideação", "7"),
        ("5.", "Seleção e Refinamento", "9"),
        ("6.", "A Solução: OnboardDev", "11"),
        ("7.", "Prototipagem", "13"),
        ("8.", "Conclusão", "14"),
        ("", "Referências", "15"),
    ]
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num} {title}")
        run.bold = True if num else False
        run.font.size = Pt(11)
        # Adicionar pontilhado + número de página
        run2 = p.add_run(f"  {'.' * (60 - len(num + title))}  {page}")
        run2.font.size = Pt(11)
        run2.font.color.rgb = GRAY_COLOR

    doc.add_page_break()

    # ================================================================
    # 1. INTRODUÇÃO
    # ================================================================
    add_heading_styled(doc, "1. Introdução", level=1)

    add_heading_styled(doc, "1.1. Contextualização", level=2)
    add_paragraph_styled(doc, (
        "A indústria de software é um dos setores de maior crescimento e transformação da economia global. "
        "Com a digitalização acelerada de empresas em todos os setores, a demanda por profissionais de "
        "tecnologia nunca esteve tão alta. No entanto, este crescimento traz consigo desafios operacionais "
        "significativos — entre eles, o processo de integração de novos desenvolvedores em projetos existentes, "
        "conhecido como onboarding."
    ))
    add_paragraph_styled(doc, (
        "O onboarding de desenvolvedores é um processo crítico que impacta diretamente a produtividade do time, "
        "a retenção de talentos e os custos operacionais da empresa. Apesar de sua importância, a maioria das "
        "organizações de software aborda este processo de forma ad hoc, desestruturada e altamente dependente "
        "de recursos humanos individuais."
    ))

    add_heading_styled(doc, "1.2. Objetivo do Trabalho", level=2)
    add_paragraph_styled(doc, (
        "Este trabalho tem como objetivo aplicar metodologias de geração de ideias — especificamente o Design "
        "Thinking — para identificar, explorar e propor uma solução inovadora para um problema real da indústria "
        "de software. O resultado é o OnboardDev, uma plataforma inteligente de onboarding para desenvolvedores, "
        "acompanhada de um protótipo funcional desenvolvido em React."
    ))

    add_heading_styled(doc, "1.3. Estrutura do Relatório", level=2)
    add_paragraph_styled(doc, (
        "O relatório segue as 6 etapas de Design Thinking propostas pela disciplina: (1) Identificação do "
        "Problema, (2) Pesquisa e Empatia, (3) Sessão de Ideação, (4) Seleção e Refinamento, (5) Descrição "
        "da Solução e (6) Prototipagem."
    ))

    # ================================================================
    # 2. IDENTIFICAÇÃO DO PROBLEMA
    # ================================================================
    add_heading_styled(doc, "2. Identificação do Problema", level=1)

    add_heading_styled(doc, "2.1. O Problema", level=2)
    add_paragraph_styled(doc, (
        "Quando um novo desenvolvedor ingressa em um time de software, ele enfrenta um período médio de 4 a 8 "
        "semanas até atingir produtividade plena. Este período é marcado por documentação desatualizada, "
        "conhecimento fragmentado entre membros do time, sobrecarga dos desenvolvedores seniores responsáveis "
        "pela mentoria informal, e ausência total de métricas objetivas de progresso."
    ))

    add_heading_styled(doc, "2.2. Impacto", level=2)
    add_paragraph_styled(doc, "O problema gera consequências mensuráveis em múltiplas dimensões:")

    create_styled_table(doc,
        headers=["Aspecto", "Impacto"],
        rows=[
            ["Custo Financeiro", "20-30% do salário anual de um dev em custos de onboarding"],
            ["Produtividade do Time", "Queda de 15-25% durante o onboarding de um novo membro"],
            ["Retenção de Talentos", "33% dos novos contratados buscam outro emprego nos primeiros 6 meses com onboarding negativo"],
            ["Time-to-Productivity", "Média de 4-8 semanas para atingir produtividade plena"],
            ["Escala", "Problema se multiplica em empresas com alta rotatividade (13,2%/ano em tech)"],
        ],
        col_widths=[5, 11]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "2.3. Justificativa", level=2)
    add_paragraph_styled(doc, (
        "O problema é universal (afeta empresas de todos os portes), recorrente (cada nova contratação o reativa), "
        "custoso (impacto direto em produtividade e retenção) e solucionável com tecnologias atuais. A ausência "
        "de soluções especializadas no mercado representa uma oportunidade significativa."
    ))

    # ================================================================
    # 3. PESQUISA E EMPATIA
    # ================================================================
    add_heading_styled(doc, "3. Pesquisa e Empatia", level=1)

    add_heading_styled(doc, "3.1. Análise Documental", level=2)
    add_paragraph_styled(doc, (
        "Foram analisadas pesquisas de mercado e relatórios da indústria para fundamentar a relevância do problema. "
        "Os dados confirmam que times com onboarding estruturado apresentam resultados significativamente superiores."
    ))

    create_styled_table(doc,
        headers=["Fonte", "Dado Relevante"],
        rows=[
            ["State of DevOps Report (DORA/Google)", "Times com bom onboarding: 2x mais deploy frequency"],
            ["Stack Overflow Developer Survey 2025", "72% consideram documentação fator decisivo para permanecer"],
            ["Glassdoor Research", "Onboarding estruturado melhora retenção em 82%"],
            ["LinkedIn Workforce Report", "Rotatividade em tech: 13,2%/ano — mais alta de qualquer indústria"],
            ["GitLab Remote Work Report", "55% dos devs remotos relatam dificuldades extremas no onboarding"],
        ],
        col_widths=[6, 10]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "3.2. Personas", level=2)
    add_paragraph_styled(doc, "Foram construídas 3 personas representando os principais stakeholders do problema:")

    # Persona 1
    add_paragraph_styled(doc, "", space_after=2)
    p = doc.add_paragraph()
    run = p.add_run("Persona 1 — Lucas Silva (Dev Junior, 23 anos): ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "Recém-contratado em uma startup de fintech com um monorepo de 200k linhas. "
        "Sente-se perdido no código, tem medo de perguntar demais. "
        "\"Eu olho pra esse código e não faço ideia por onde começar.\""
    )
    run.font.size = Pt(10)
    run.italic = True

    # Persona 2
    p = doc.add_paragraph()
    run = p.add_run("Persona 2 — Marina Costa (Dev Sênior, 31 anos): ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "Pessoa de referência do time, interrompida 5-10 vezes por dia com perguntas de novos devs. "
        "\"Eu amo ajudar, mas não consigo entregar minhas próprias tarefas.\""
    )
    run.font.size = Pt(10)
    run.italic = True

    # Persona 3
    p = doc.add_paragraph()
    run = p.add_run("Persona 3 — Ricardo Mendes (Tech Lead, 35 anos): ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "Gerencia 8 devs e contratou 2 juniors. Não tem visibilidade real do progresso. "
        "\"Eu preciso de números, não de sensações.\""
    )
    run.font.size = Pt(10)
    run.italic = True

    add_heading_styled(doc, "3.3. Mapa de Empatia — Lucas (Dev Junior)", level=2)
    create_styled_table(doc,
        headers=["Dimensão", "Conteúdo"],
        rows=[
            ["Pensa e Sente", "Medo de parecer incompetente, sobrecarga de informação, ansiedade por provar competência"],
            ["Ouve", "\"Leia a documentação\", \"Pergunta pro fulano\", \"Todo mundo passa por isso\""],
            ["Vê", "Repositório com centenas de arquivos, colegas produtivos, documentação desatualizada"],
            ["Fala e Faz", "Faz perguntas básicas, lê código tentando entender, anota descobertas"],
            ["Dores", "Falta de direção, medo de errar, informação fragmentada, feedback tardio"],
            ["Ganhos", "Primeiro PR mergeado, elogio do mentor, entender um fluxo sozinho"],
        ],
        col_widths=[4, 12]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "3.4. Jornada do Usuário", level=2)
    add_paragraph_styled(doc, (
        "Foi mapeada a jornada de onboarding na situação atual (sem OnboardDev), evidenciando 7 fases desde o "
        "primeiro dia até a produtividade plena (semana 7-8), com emoções variando de ansiedade a alívio. "
        "Em contraste, a jornada ideal com OnboardDev reduz o processo para 2-3 semanas com emoções "
        "predominantemente positivas."
    ))

    create_styled_table(doc,
        headers=["Fase", "Sem OnboardDev", "Com OnboardDev"],
        rows=[
            ["Dia 1", "Perdido, sem guia estruturado 😰", "Mapa interativo, visão macro 😃"],
            ["Dia 2-3", "README desatualizado 😤", "Trilha guiada com progresso 😊"],
            ["Semana 1", "Dependência do sênior 😵", "Tarefas com dicas contextuais 🚀"],
            ["Semana 2", "Primeira task sem contexto 😓", "Primeiro PR com confiança 💪"],
            ["Semana 3-4", "PR rejeitado por convenções 😞", "Trilha completa, badge conquistado 🎉"],
            ["Semana 7-8", "Finalmente produtivo 😐", "Produtivo desde a semana 2 ✅"],
        ],
        col_widths=[3, 6.5, 6.5]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "3.5. Insights Principais", level=2)
    insights = [
        "A solução precisa ser visual e interativa — documentação estática não funciona",
        "Autonomia é essencial — o dev quer aprender no seu ritmo, sem dependência",
        "Métricas objetivas são necessárias para o gestor tomar decisões informadas",
        "A solução deve ser low-maintenance para não ficar desatualizada como docs manuais",
        "Gamificação leve aumenta engajamento do público-alvo (devs jovens habituados a interfaces gamificadas)",
    ]
    for i, insight in enumerate(insights, 1):
        add_bullet(doc, insight)

    # ================================================================
    # 4. SESSÃO DE IDEAÇÃO
    # ================================================================
    add_heading_styled(doc, "4. Sessão de Ideação", level=1)

    add_heading_styled(doc, "4.1. Metodologia", level=2)
    add_paragraph_styled(doc, (
        "A sessão de ideação foi conduzida em 90 minutos utilizando três técnicas complementares: "
        "How Might We (HMW) para reframe dos problemas, Brainstorming Livre para geração divergente, "
        "e Crazy 8s para ideação visual rápida."
    ))

    add_heading_styled(doc, "4.2. Perguntas How Might We", level=2)
    hmw_items = [
        "Como poderíamos fazer um dev entender a arquitetura de um projeto em minutos, não semanas?",
        "Como poderíamos eliminar a dependência de devs seniores para o onboarding?",
        "Como poderíamos tornar a documentação auto-atualizável?",
        "Como poderíamos dar ao Tech Lead visibilidade em tempo real do progresso?",
        "Como poderíamos transformar o onboarding em uma experiência prazerosa?",
        "Como poderíamos garantir que o novo dev aprenda as convenções antes de errar?",
        "Como poderíamos escalar o onboarding sem escalar o custo de mentoria?",
    ]
    for item in hmw_items:
        add_bullet(doc, item)

    add_heading_styled(doc, "4.3. Resultados do Brainstorming", level=2)
    add_paragraph_styled(doc, "Foram geradas 25 ideias organizadas em 5 categorias:")

    create_styled_table(doc,
        headers=["Categoria", "Qty", "Exemplos de Ideias"],
        rows=[
            ["Visualização do Código", "5", "Mapa interativo, Code Explorer, Heatmap de complexidade"],
            ["Aprendizado Guiado", "5", "Trilhas progressivas, Tasks práticas, Playlists de conteúdo"],
            ["Documentação Inteligente", "5", "Docs por IA, Changelog narrativo, Chat com codebase"],
            ["Métricas e Gestão", "5", "Dashboard de progresso, Alertas de bloqueio, Painel do time"],
            ["Gamificação", "5", "Badges, Streaks, Leaderboard, Celebrações visuais"],
        ],
        col_widths=[5, 1.5, 9.5]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "4.4. Análise SWOT", level=2)
    add_paragraph_styled(doc, "Uma análise SWOT completa foi conduzida para avaliar a viabilidade estratégica do OnboardDev:")

    create_styled_table(doc,
        headers=["", "Positivo", "Negativo"],
        rows=[
            ["Interno", "FORÇAS: Proposta clara, interface visual inovadora, métricas inéditas no mercado, gamificação, ROI demonstrável",
             "FRAQUEZAS: Dependência de IA, necessidade multi-plataforma, configuração inicial complexa, equipe pequena"],
            ["Externo", "OPORTUNIDADES: Mercado de DevTools US$ 32B, trabalho remoto crescente, LLMs acessíveis, nicho pouco explorado",
             "AMEAÇAS: Grandes players (GitHub/GitLab), IA genérica (Copilot), resistência organizacional, segurança de código"],
        ],
        col_widths=[3, 6.5, 6.5]
    )

    doc.add_paragraph("")
    add_paragraph_styled(doc, (
        "Estratégias cruzadas SO, WO, ST e WT foram derivadas. A principal conclusão é que o OnboardDev ocupa "
        "uma posição estratégica favorável com diferenciais de nicho que soluções genéricas não oferecem."
    ))

    add_heading_styled(doc, "4.5. Mapa Mental", level=2)
    add_paragraph_styled(doc, (
        "Foi construído um mapa mental abrangente cobrindo 6 dimensões do OnboardDev: Problema, Público-Alvo, "
        "Funcionalidades Core (6 módulos), Tecnologias (React + TypeScript), Modelo de Negócio (SaaS B2B com "
        "3 tiers) e Diferenciais Competitivos. O mapa revelou conexões importantes entre os ramos, como a "
        "relação direta entre cada funcionalidade e a dor específica que resolve."
    ))

    # ================================================================
    # 5. SELEÇÃO E REFINAMENTO
    # ================================================================
    add_heading_styled(doc, "5. Seleção e Refinamento", level=1)

    add_heading_styled(doc, "5.1. Critérios de Seleção", level=2)
    add_paragraph_styled(doc, "As 25 ideias foram avaliadas em 5 critérios ponderados:")

    create_styled_table(doc,
        headers=["Critério", "Peso", "Descrição"],
        rows=[
            ["Impacto no Problema", "30%", "O quanto resolve diretamente a dor do onboarding"],
            ["Viabilidade Técnica", "25%", "Possibilidade de implementar como protótipo no prazo"],
            ["Originalidade", "20%", "Inovação em relação a soluções existentes"],
            ["Experiência do Usuário", "15%", "Melhoria na experiência do dev durante o onboarding"],
            ["Escalabilidade", "10%", "Capacidade de escalar para diferentes times/projetos"],
        ],
        col_widths=[5, 2, 9]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "5.2. Matriz de Priorização", level=2)
    add_paragraph_styled(doc, (
        "As ideias foram classificadas em quadrantes de Impacto × Esforço. As ideias de alto impacto e baixo "
        "esforço (Fazer Primeiro) foram: Mapa Interativo, Trilhas Progressivas, Dashboard de Progresso, "
        "Badges e Tasks Guiadas."
    ))

    add_heading_styled(doc, "5.3. Ranking das Top 10 Ideias", level=2)
    create_styled_table(doc,
        headers=["#", "Ideia", "Score"],
        rows=[
            ["1", "Trilhas de Onboarding Progressivas", "9.05"],
            ["2", "Mapa Interativo da Arquitetura", "8.95"],
            ["3", "Dashboard de Progresso", "8.50"],
            ["4", "Tasks Práticas Guiadas", "8.35"],
            ["5", "Code Explorer com Explicações IA", "8.30"],
            ["6", "Painel de Gestão do Time", "7.85"],
            ["7", "Documentação Gerada por IA", "7.65"],
            ["8", "Badges e Conquistas", "7.15"],
            ["9", "Timeline de Atividades", "6.70"],
            ["10", "Streak de Dias Consecutivos", "6.65"],
        ],
        col_widths=[1.5, 9, 2.5]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "5.4. Funcionalidades Selecionadas", level=2)
    add_paragraph_styled(doc, "As funcionalidades foram organizadas em 3 tiers de prioridade:")

    add_paragraph_styled(doc, "", space_after=2)
    p = doc.add_paragraph()
    run = p.add_run("Tier 1 (MVP): ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Trilhas de Onboarding, Mapa Interativo, Dashboard de Progresso, Tasks Guiadas, Landing Page + Auth")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Tier 2 (Importante): ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Code Explorer com IA, Painel de Gestão do Time, Documentação por IA")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Tier 3 (Nice-to-have): ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Badges, Streaks, Timeline de Atividades, Configurações")
    run.font.size = Pt(10)

    # ================================================================
    # 6. A SOLUÇÃO: ONBOARDDEV
    # ================================================================
    add_heading_styled(doc, "6. A Solução: OnboardDev", level=1)

    add_heading_styled(doc, "6.1. Visão do Produto", level=2)
    add_paragraph_styled(doc, (
        "O OnboardDev é uma plataforma web inteligente composta por 11 módulos integrados que cobrem todo o "
        "ciclo de onboarding — desde a primeira visualização da arquitetura até a confirmação de produtividade "
        "plena pelo Tech Lead. A proposta de valor central é: \"Domine qualquer codebase em dias, não semanas.\""
    ))

    add_heading_styled(doc, "6.2. Módulos do Sistema", level=2)
    create_styled_table(doc,
        headers=["#", "Módulo", "Descrição"],
        rows=[
            ["1", "Landing Page", "Apresentação do produto com hero, features e CTA"],
            ["2", "Autenticação", "Login simulado com modo demo para exploração"],
            ["3", "Dashboard", "Hub central com repositórios, métricas e atividades recentes"],
            ["4", "Mapa Interativo", "Diagrama visual navegável com módulos e conexões (diferencial)"],
            ["5", "Trilhas de Onboarding", "Caminho estruturado com tarefas progressivas"],
            ["6", "Detalhe da Tarefa", "Instruções, código de referência e objetivo de aprendizado"],
            ["7", "Progresso do Dev", "Gráficos, métricas, badges e heatmap de atividade"],
            ["8", "Painel do Time", "Visão do Tech Lead com alertas e métricas agregadas"],
            ["9", "Docs por IA", "Documentação auto-gerada com API reference e exemplos"],
            ["10", "Code Explorer", "Navegação pelo código com explicações contextuais"],
            ["11", "Configurações", "Personalização da experiência de onboarding"],
        ],
        col_widths=[1, 4, 11]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "6.3. Fluxos de Uso Principais", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Fluxo do Dev Junior: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Landing → Demo → Dashboard → Mapa → Trilha → Tarefas → Progresso → Badges")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Fluxo do Tech Lead: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Login → Dashboard → Painel do Time → Identificar bloqueio → Ação")
    run.font.size = Pt(10)

    add_heading_styled(doc, "6.4. Diferenciais Competitivos", level=2)
    create_styled_table(doc,
        headers=["Aspecto", "Soluções Atuais", "OnboardDev"],
        rows=[
            ["Visualização", "Documentação textual estática", "Mapa interativo navegável"],
            ["Aprendizado", "\"Leia o README\"", "Trilhas progressivas com tarefas"],
            ["Documentação", "Manual, desatualiza rápido", "Auto-gerada por IA, sempre atual"],
            ["Progresso", "\"Acho que tá indo bem\"", "Métricas objetivas com gráficos"],
            ["Mentoria", "Dependente de 1-2 seniores", "Automatizada e escalável"],
            ["Engajamento", "Nenhum", "Badges, streaks, celebrações"],
        ],
        col_widths=[3.5, 5.5, 7]
    )

    # ================================================================
    # 7. PROTOTIPAGEM
    # ================================================================
    doc.add_paragraph("")
    add_heading_styled(doc, "7. Prototipagem", level=1)

    add_heading_styled(doc, "7.1. Abordagem", level=2)
    add_paragraph_styled(doc, (
        "O protótipo foi desenvolvido em duas fases: (1) geração inicial no Google Stitch para criação rápida "
        "das telas base e (2) refinamento e desenvolvimento completo local em React + TypeScript com Vite."
    ))

    add_heading_styled(doc, "7.2. Tecnologias Utilizadas", level=2)
    create_styled_table(doc,
        headers=["Tecnologia", "Versão", "Uso"],
        rows=[
            ["React", "18+", "Framework de UI com componentes reutilizáveis"],
            ["TypeScript", "5+", "Tipagem estática para segurança e produtividade"],
            ["Vite", "5+", "Build tool para desenvolvimento rápido"],
            ["React Router", "v6", "Navegação SPA entre páginas"],
            ["Recharts", "2+", "Gráficos e visualizações de dados"],
            ["Lucide React", "—", "Biblioteca de ícones"],
            ["CSS Modules", "—", "Estilização modular e escopada"],
            ["localStorage", "—", "Persistência de dados mockados"],
        ],
        col_widths=[4, 2, 10]
    )

    doc.add_paragraph("")
    add_heading_styled(doc, "7.3. Telas Implementadas", level=2)
    add_paragraph_styled(doc, (
        "O protótipo inclui todas as 12 telas planejadas com navegação funcional, dados mockados realistas, "
        "design premium com dark mode, micro-animações e responsividade. O sistema é completamente funcional "
        "como uma Single Page Application (SPA), com estado persistido em localStorage."
    ))

    add_heading_styled(doc, "7.4. Dados de Demonstração", level=2)
    add_paragraph_styled(doc, (
        "O protótipo vem pré-carregado com dados realistas simulando um projeto de e-commerce, incluindo: "
        "3 repositórios conectados, 4 trilhas de onboarding com 26 tarefas progressivas, 10 módulos de "
        "arquitetura com conexões, 4 desenvolvedores em diferentes estágios de onboarding, 8 badges de "
        "conquista, e documentação auto-gerada de 5 módulos."
    ))

    # ================================================================
    # 8. CONCLUSÃO
    # ================================================================
    add_heading_styled(doc, "8. Conclusão", level=1)

    add_heading_styled(doc, "8.1. Resultados Alcançados", level=2)
    add_paragraph_styled(doc, (
        "O processo de Design Thinking aplicado permitiu identificar um problema real e relevante na indústria "
        "de software, investigá-lo em profundidade através de pesquisa e empatia, gerar 25 ideias por meio de "
        "sessões estruturadas de ideação, selecionar as mais promissoras com critérios objetivos, e "
        "materializá-las em um protótipo funcional."
    ))
    add_paragraph_styled(doc, (
        "O OnboardDev representa uma proposta inovadora que preenche uma lacuna clara no mercado de ferramentas "
        "para desenvolvedores, oferecendo uma experiência integrada de onboarding que nenhuma solução existente "
        "oferece atualmente."
    ))

    add_heading_styled(doc, "8.2. Lições Aprendidas", level=2)
    lessons = [
        "A empatia direciona a inovação — as personas e jornadas foram fundamentais para priorizar funcionalidades",
        "Quantidade gera qualidade — as 25 ideias do brainstorming geraram insights que refinaram a solução",
        "Critérios objetivos evitam viés — a matriz de priorização com scoring evitou preferências pessoais",
        "Prototipar é validar — a construção revelou nuances invisíveis na especificação teórica",
    ]
    for lesson in lessons:
        add_bullet(doc, lesson)

    add_heading_styled(doc, "8.3. Próximos Passos", level=2)
    next_steps = [
        "Desenvolvimento de backend com integração real à API do GitHub",
        "Implementação de análise de código com Large Language Models (LLMs)",
        "Validação com usuários reais em empresas parceiras",
        "Modelo de monetização SaaS com tier freemium",
    ]
    for step in next_steps:
        add_bullet(doc, step)

    # ================================================================
    # REFERÊNCIAS
    # ================================================================
    add_heading_styled(doc, "Referências", level=1)

    references = [
        "DORA Team / Google Cloud. Accelerate State of DevOps Report 2025. Google, 2025.",
        "Stack Overflow. Developer Survey 2025. Stack Overflow, 2025.",
        "Glassdoor. The True Cost of a Bad Hire. Glassdoor Economic Research, 2024.",
        "LinkedIn. Workforce Report: Tech Industry Turnover Trends. LinkedIn, 2025.",
        "Bureau of Labor Statistics. Employee Tenure Summary. U.S. Department of Labor, 2025.",
        "Harvard Business Review. Getting New Hires Up to Speed Quickly. HBR, 2024.",
        "GitLab. The Remote Work Report 2025. GitLab Inc., 2025.",
        "Brown, T. Design Thinking. Harvard Business Review, 2008.",
        "Ries, E. The Lean Startup. Crown Business, 2011.",
        "Norman, D. The Design of Everyday Things. Basic Books, 2013.",
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(ref)
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(2)

    # ================================================================
    # Salvar
    # ================================================================
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\n[OK] Relatorio gerado com sucesso!")
    print(f"Arquivo: {OUTPUT_FILE}")
    print(f"Total de paginas estimado: ~15")


if __name__ == "__main__":
    generate_report()
